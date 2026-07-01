from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[2]
OUTPUT_DIR = ROOT / "output" / "doc"
BACKUP_DIR = ROOT / "tmp" / "docs" / "before_station_monitoring"

FUNCTIONAL_REPORT = (
    OUTPUT_DIR
    / "01_Informe_de_Alcance_y_Diseno_Funcional_de_la_Plataforma_de_Visualizacion_Sismica.docx"
)
TECHNICAL_REPORT = (
    OUTPUT_DIR
    / "02_Informe_Tecnico_de_Arquitectura_Desarrollo_y_Entorno_WSL2_de_la_Plataforma_de_Visualizacion_Sismica.docx"
)

MOJIBAKE_REPLACEMENTS = {
    "Versi?n": "Versión",
    "ejecuci?n": "ejecución",
    "variar?n": "variarán",
    "s?smico": "sísmico",
    "intensidad m?xima": "intensidad máxima",
    "clave ?nica": "clave única",
    "?nico": "único",
    "actualizaci?n": "actualización",
    "correcci?n": "corrección",
    "Jap?n": "Japón",
    "can?nico": "canónico",
    "referencias ?nicas": "referencias únicas",
    "evidencia espec?fica": "evidencia específica",
    "Exposici?n": "Exposición",
}


def replace_text_preserving_style(paragraph, replacements: dict[str, str]) -> None:
    original = paragraph.text
    updated = original
    for old, new in replacements.items():
        updated = updated.replace(old, new)
    if updated == original:
        return

    if not paragraph.runs:
        paragraph.text = updated
        return

    first = paragraph.runs[0]
    first.text = updated
    for run in paragraph.runs[1:]:
        run.text = ""


def replace_mojibake(document: DocumentObject) -> None:
    for paragraph in document.paragraphs:
        replace_text_preserving_style(paragraph, MOJIBAKE_REPLACEMENTS)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_text_preserving_style(paragraph, MOJIBAKE_REPLACEMENTS)


def set_cell_shading(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def format_table(table, widths_cm: list[float] | None = None) -> None:
    table.style = "Table Grid"
    table.autofit = True
    for cell in table.rows[0].cells:
        set_cell_shading(cell, "D9E2F3")
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
    for row in table.rows[1:]:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
    if widths_cm:
        for row in table.rows:
            for index, width in enumerate(widths_cm):
                if index < len(row.cells):
                    row.cells[index].width = Cm(width)


def add_table(document: DocumentObject, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    for index, value in enumerate(headers):
        table.rows[0].cells[index].text = value
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = value
    format_table(table, widths)
    document.add_paragraph()


def insert_paragraph_after(paragraph, text: str, style: str):
    new_element = OxmlElement("w:p")
    paragraph._p.addnext(new_element)
    new_paragraph = paragraph._parent.add_paragraph()
    new_paragraph._p.getparent().remove(new_paragraph._p)
    new_element.addnext(new_paragraph._p)
    new_paragraph.style = style
    new_paragraph.add_run(text)
    return new_paragraph


def ensure_index_entry(document: DocumentObject, previous_entry: str, new_entry: str) -> None:
    if any(p.style.name == "List Bullet" and p.text.strip() == new_entry for p in document.paragraphs):
        return
    for paragraph in document.paragraphs:
        if paragraph.style.name == "List Bullet" and paragraph.text.strip() == previous_entry:
            insert_paragraph_after(paragraph, new_entry, "List Bullet")
            return
    raise RuntimeError(f"No se encontro la entrada de indice: {previous_entry}")


def update_cover_version(document: DocumentObject, version: str) -> None:
    for paragraph in document.paragraphs[:12]:
        if paragraph.text.startswith("Versión "):
            replace_text_preserving_style(
                paragraph,
                {paragraph.text: f"Versión {version} - Junio de 2026"},
            )
            return
    raise RuntimeError("No se encontro la version en la portada")


def update_change_control(document: DocumentObject, text: str) -> None:
    for table in document.tables[:3]:
        for row in table.rows:
            if row.cells[0].text.strip() == "Control de cambios":
                row.cells[1].text = text
                return
    raise RuntimeError("No se encontro la fila Control de cambios")


def add_bullet(document: DocumentObject, text: str) -> None:
    document.add_paragraph(text, style="List Bullet")


def add_number(document: DocumentObject, text: str) -> None:
    document.add_paragraph(text, style="List Number")


def add_page_break(document: DocumentObject) -> None:
    paragraph = document.add_paragraph()
    paragraph.add_run().add_break(WD_BREAK.PAGE)


def normalize_section_lists(document: DocumentObject, heading_prefix: str) -> None:
    in_section = False
    for paragraph in document.paragraphs:
        if paragraph.style.name == "Heading 1":
            if paragraph.text.strip().startswith(heading_prefix):
                in_section = True
                continue
            if in_section:
                break
        if in_section and paragraph.style.name == "List Number":
            paragraph.style = "List Bullet"


def remove_page_break_before_heading(document: DocumentObject, heading_prefix: str) -> None:
    for paragraph in document.paragraphs:
        if paragraph.style.name != "Heading 1" or not paragraph.text.strip().startswith(heading_prefix):
            continue
        previous = paragraph._p.getprevious()
        if previous is not None and previous.xpath(".//w:br[@w:type='page']"):
            previous.getparent().remove(previous)
        return


def set_page_break_before_heading(document: DocumentObject, heading_text: str) -> None:
    for paragraph in document.paragraphs:
        if paragraph.style.name == "Heading 2" and paragraph.text.strip() == heading_text:
            paragraph.paragraph_format.page_break_before = True
            paragraph.paragraph_format.keep_with_next = True
            return


def append_functional_section(document: DocumentObject) -> None:
    if any(
        p.style.name == "Heading 1"
        and p.text.strip().startswith("18. Subsistema experimental")
        for p in document.paragraphs
    ):
        return

    add_page_break(document)
    document.add_heading(
        "18. Subsistema experimental de estaciones y propagación sísmica",
        level=1,
    )
    document.add_paragraph(
        "La plataforma incorpora un subsistema experimental orientado a visualizar estaciones "
        "sismológicas, representar su condición operativa y recibir resultados calculados por "
        "un motor científico externo. El propósito es aproximar la experiencia funcional de un "
        "monitor de formas de onda sin sustituir la información oficial del IGP/CENSIS, USGS, "
        "CSN u otras autoridades."
    )

    document.add_heading("18.1 Alcance funcional aprobado", level=2)
    add_table(
        document,
        ["Componente", "Incluido", "Responsabilidad"],
        [
            [
                "Catálogo de estaciones",
                "Sí",
                "Obtener metadatos FDSN, ubicación, red, periodo operativo y atribución.",
            ],
            [
                "Estado de estación",
                "Sí",
                "Mostrar estado desconocido, en línea, con retraso, fuera de línea o activado.",
            ],
            [
                "Propagación P/S",
                "Sí",
                "Representar frentes estimados con tiempo de origen, profundidad y velocidades configuradas.",
            ],
            [
                "Recepción científica",
                "Sí, mediante adaptador",
                "Aceptar estados, picks y orígenes experimentales producidos por SeisComP.",
            ],
            [
                "Detección propia en navegador",
                "No",
                "El navegador no procesa MiniSEED ni decide si ocurrió un sismo.",
            ],
            [
                "Alerta oficial",
                "No",
                "Los resultados experimentales no reemplazan ni retrasan publicaciones institucionales.",
            ],
        ],
        [4.3, 3.0, 9.0],
    )

    document.add_heading("18.2 Interacción y representación", level=2)
    add_bullet(
        document,
        "Las estaciones se representan mediante símbolos triangulares estables sobre el globo 3D.",
    )
    add_bullet(
        document,
        "El color indica estado operativo o activación, no una intensidad sísmica oficial.",
    )
    add_bullet(
        document,
        "La selección de una estación presenta red, código, ubicación, latencia y última observación.",
    )
    add_bullet(
        document,
        "Los frentes P y S se sincronizan con la hora de origen del evento y se identifican como estimaciones.",
    )
    add_bullet(
        document,
        "El usuario puede ocultar estaciones para preservar legibilidad y rendimiento del mapa.",
    )

    document.add_heading("18.3 Principios de uso responsable", level=2)
    add_number(
        document,
        "Mantener separación explícita entre eventos oficiales y orígenes experimentales.",
    )
    add_number(
        document,
        "No publicar una advertencia de evacuación, tsunami o daño a partir del motor experimental.",
    )
    add_number(
        document,
        "Conservar fuente, estación, instante UTC, latencia, algoritmo y versión para auditoría.",
    )
    add_number(
        document,
        "Operar el procesamiento científico en Linux con software especializado; Node.js integra y distribuye resultados.",
    )
    add_number(
        document,
        "Validar por reproducción histórica antes de habilitar procesamiento continuo.",
    )
    add_number(
        document,
        "No copiar código ni interfaz de GlobalQuake 1.1.2; su desarrollo posterior a 1.0 es propietario.",
    )

    document.add_heading("18.4 Fuentes y dependencia institucional", level=2)
    document.add_paragraph(
        "GEOFON ofrece metadatos FDSN y un servicio SeedLink público para estaciones abiertas. "
        "EarthScope distribuye datos abiertos en tiempo casi real y metadatos StationXML. El IGP "
        "confirma transmisión en tiempo real dentro de la Red Sísmica Nacional, pero no publica "
        "un endpoint SeedLink documentado para consumo general; una integración directa con "
        "estaciones peruanas requiere coordinación institucional."
    )

    document.add_heading("18.5 Trazabilidad", level=2)
    document.add_paragraph(
        "La especificación técnica vinculante es SDD-014. Las pruebas previstas se registran en "
        "TEST-014 y la evidencia ejecutada en VALIDATION-014. El módulo permanece identificado "
        "como experimental hasta superar validación histórica, revisión sismológica y evaluación operativa."
    )

    document.add_heading("18.6 Referencias", level=2)
    add_bullet(document, "GEOFON SeedLink: https://geofon.gfz.de/waveform/seedlink.php")
    add_bullet(
        document,
        "GEOFON FDSN Station: https://geofon.gfz.de/fdsnws/station/1/",
    )
    add_bullet(
        document,
        "FDSN StationXML: https://www.fdsn.org/xml/station/",
    )
    add_bullet(
        document,
        "SeisComP scautopick: https://docs.gempa.de/seiscomp/current/apps/scautopick.html",
    )
    add_bullet(
        document,
        "SeisComP scautoloc: https://docs.gempa.de/seiscomp/current/apps/scautoloc.html",
    )
    add_bullet(
        document,
        "IGP Red Sísmica Nacional: https://ultimosismo.igp.gob.pe/red-sismica-nacional",
    )


def append_technical_section(document: DocumentObject) -> None:
    if any(
        p.style.name == "Heading 1"
        and p.text.strip().startswith("17. Arquitectura del subsistema")
        for p in document.paragraphs
    ):
        return

    add_page_break(document)
    document.add_heading(
        "17. Arquitectura del subsistema experimental de estaciones",
        level=1,
    )
    document.add_paragraph(
        "La solución adopta una arquitectura desacoplada. SeisComP realiza adquisición, picking, "
        "asociación, localización y magnitud en Linux. TypeScript y Node.js conservan la "
        "responsabilidad de normalizar contratos, persistir resultados, aplicar seguridad y "
        "distribuir estados hacia CesiumJS."
    )

    document.add_heading("17.1 Flujo técnico", level=2)
    flow = document.add_paragraph(style="CodeBlock")
    flow.add_run(
        "FDSN Station / SeedLink\n"
        "          |\n"
        "          v\n"
        "SeisComP en Linux VM o WSL2\n"
        "  scautopick -> scautoloc -> scamp/scmag\n"
        "          |\n"
        "          v\n"
        "Adaptador interno autenticado\n"
        "          |\n"
        "          v\n"
        "Node.js API/worker -> PostgreSQL/PostGIS -> SSE -> CesiumJS"
    )

    document.add_heading("17.2 Componentes", level=2)
    add_table(
        document,
        ["Componente", "Tecnología", "Responsabilidad"],
        [
            [
                "Station catalog provider",
                "TypeScript + FDSN text",
                "Descargar, validar y versionar metadatos de estaciones.",
            ],
            [
                "Scientific engine",
                "SeisComP sobre Linux",
                "Consumir SeedLink/MiniSEED y producir picks, amplitudes y orígenes.",
            ],
            [
                "Engine adapter",
                "API interna Node.js",
                "Validar autenticación, esquema, límites temporales e idempotencia.",
            ],
            [
                "Persistence",
                "PostgreSQL + PostGIS",
                "Guardar estaciones, último estado, picks y orígenes experimentales.",
            ],
            [
                "Distribution",
                "REST + SSE",
                "Exponer catálogo y cambios sin abrir SeedLink al navegador.",
            ],
            [
                "Visualization",
                "React + CesiumJS",
                "Renderizar estaciones, detalle y frentes P/S sincronizados.",
            ],
        ],
        [4.0, 4.2, 8.2],
    )

    document.add_heading("17.3 Modelo de datos", level=2)
    add_table(
        document,
        ["Entidad", "Clave", "Contenido principal"],
        [
            [
                "seismic_stations",
                "station_id",
                "Red, estación, coordenadas, elevación, sitio, fechas y fuente.",
            ],
            [
                "station_states",
                "station_id",
                "Estado vigente, latencia, fase, valor de activación y hora UTC.",
            ],
            [
                "seismic_picks",
                "pick_id",
                "Estación, fase, tiempo de llegada, SNR, amplitud y algoritmo.",
            ],
            [
                "experimental_origins",
                "origin_id",
                "Hipocentro, magnitud, calidad, estaciones usadas y estado de revisión.",
            ],
        ],
        [4.6, 4.1, 7.7],
    )

    document.add_heading("17.4 Contratos API", level=2)
    add_table(
        document,
        ["Método y ruta", "Uso", "Exposición"],
        [
            ["GET /api/stations", "Catálogo y último estado", "Pública, solo lectura"],
            ["GET /api/stations/stream", "Cambios de estado por SSE", "Pública, solo lectura"],
            [
                "POST /internal/seismic-engine/snapshots",
                "Estados y picks producidos por el motor",
                "Interna, token obligatorio",
            ],
            [
                "POST /internal/seismic-engine/origins",
                "Orígenes experimentales",
                "Interna, token obligatorio",
            ],
        ],
        [4.8, 6.5, 5.1],
    )

    document.add_heading("17.5 Reglas de seguridad y calidad", level=2)
    add_bullet(
        document,
        "El token del adaptador se mantiene fuera del repositorio y nunca se entrega al frontend.",
    )
    add_bullet(
        document,
        "Cada mensaje tiene identificador idempotente, marca UTC y versión de esquema.",
    )
    add_bullet(
        document,
        "Se rechazan coordenadas, fases, latencias o tiempos fuera de límites configurados.",
    )
    add_bullet(
        document,
        "Los orígenes experimentales no ingresan automáticamente al catálogo oficial consolidado.",
    )
    add_bullet(
        document,
        "La conexión SeedLink se centraliza en el motor; los navegadores no se conectan a redes científicas.",
    )

    document.add_heading("17.6 Despliegue y operación", level=2)
    document.add_paragraph(
        "Para desarrollo se admite Windows como host con Ubuntu 24.04 sobre WSL2 o una máquina "
        "virtual Linux. PostgreSQL y las aplicaciones TypeScript pueden continuar en el entorno "
        "actual. SeisComP debe instalarse y configurarse dentro de Linux con almacenamiento "
        "independiente para buffers MiniSEED. El primer modo operativo es reproducción histórica "
        "y luego modo sombra en tiempo real."
    )

    document.add_heading("17.7 Criterios técnicos de salida", level=2)
    add_number(document, "Catálogo FDSN importado sin duplicar station_id.")
    add_number(document, "API de estaciones validada por esquema y pruebas de integración.")
    add_number(document, "Estados del adaptador persistidos de forma idempotente.")
    add_number(document, "CesiumJS representa estaciones sin degradar la interacción del globo.")
    add_number(document, "Frentes P/S usan el tiempo de origen y no reinician al seleccionar.")
    add_number(document, "Pruebas automatizadas, build y validación visual aprobados.")
    add_number(document, "Toda salida experimental conserva etiqueta visible y trazabilidad.")

    document.add_heading("17.8 Trazabilidad documental", level=2)
    document.add_paragraph(
        "Esta ampliación deriva del informe funcional, del presente informe técnico y de la "
        "norma PROCESO_DE_ENTREGA_Y_VALIDACION. Su implementación se rige por SDD-014; "
        "TEST-014 define las pruebas y VALIDATION-014 registra los resultados."
    )


def prepare_backup(path: Path) -> None:
    BACKUP_DIR.mkdir(parents=True, exist_ok=True)
    backup = BACKUP_DIR / path.name
    if not backup.exists():
        backup.write_bytes(path.read_bytes())


def update_document(
    path: Path,
    *,
    version: str,
    previous_index: str,
    new_index: str,
    control_text: str,
    append_section,
    section_heading_prefix: str,
    remove_section_page_break: bool = False,
) -> None:
    prepare_backup(path)
    document = Document(path)
    replace_mojibake(document)
    update_cover_version(document, version)
    update_change_control(document, control_text)
    ensure_index_entry(document, previous_index, new_index)
    append_section(document)
    normalize_section_lists(document, section_heading_prefix)
    if remove_section_page_break:
        remove_page_break_before_heading(document, section_heading_prefix)
    if path == TECHNICAL_REPORT:
        set_page_break_before_heading(document, "17.4 Contratos API")
    document.save(path)


def main() -> None:
    update_document(
        FUNCTIONAL_REPORT,
        version="1.7",
        previous_index="17. Referencias complementarias",
        new_index="18. Subsistema experimental de estaciones y propagación sísmica",
        control_text=(
            "Versión 1.7 incorpora el alcance funcional del subsistema experimental de "
            "estaciones, propagación P/S, integración con un motor científico externo y "
            "límites explícitos respecto de alertas oficiales."
        ),
        append_section=append_functional_section,
        section_heading_prefix="18. Subsistema experimental",
    )
    update_document(
        TECHNICAL_REPORT,
        version="1.7",
        previous_index="16. Referencias complementarias",
        new_index="17. Arquitectura del subsistema experimental de estaciones",
        control_text=(
            "Versión 1.7 define la arquitectura desacoplada FDSN/SeedLink - SeisComP - "
            "Node.js - PostgreSQL/PostGIS - CesiumJS, sus contratos internos y criterios "
            "de validación."
        ),
        append_section=append_technical_section,
        section_heading_prefix="17. Arquitectura del subsistema",
        remove_section_page_break=True,
    )


if __name__ == "__main__":
    main()
